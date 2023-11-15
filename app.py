import openai
from flask import Flask, request, send_file, jsonify, make_response, send_from_directory
from flask_cors import CORS
from docx import Document
import os
from werkzeug.datastructures import FileStorage
import tempfile
import json
from werkzeug.exceptions import RequestEntityTooLarge
import logging
import traceback
import nltk
from nltk.tokenize import sent_tokenize
from concurrent.futures import ProcessPoolExecutor
from docx.shared import RGBColor
import zipfile

nltk.download('punkt')

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)
app.debug = False

@app.errorhandler(RequestEntityTooLarge)
def handle_large_payload(error):
    return 'Request payload too large', 413

uploaded_templates = {}

@app.route('/upload-template', methods=['POST'])
def upload_template():
    template_files = request.files.getlist('templates')  # Retrieve list of files
    template_type = request.form.get('templateType')

    # Check if all files are .docx format
    for template_file in template_files:
        if not template_file.filename.endswith('.docx'):
            return jsonify({"error": "Invalid file format! Only .docx is allowed."}), 400

    # Save all uploaded templates
    saved_paths = []
    for template_file in template_files:
        temp_path = tempfile.mktemp(suffix=".docx")
        template_file.save(temp_path)
        saved_paths.append(temp_path)

    # Append or set the saved paths to uploaded_templates
    if template_type in uploaded_templates:
        uploaded_templates[template_type].extend(saved_paths)
    else:
        uploaded_templates[template_type] = saved_paths
    
    uploaded_filenames = [template_file.filename for template_file in template_files]

    return jsonify({"message": "Templates uploaded successfully!", "uploadedFiles": uploaded_filenames})

def get_answer(content, question, model="gpt-4"):
    try:
        messages = [
            {"role": "system", "content": "You are an expert in legal document processing, analysis, and modification. You understand agreements, clauses, and contract details thoroughly. Help users with their document-related questions and modifications efficiently and accurately."},
            {"role": "user", "content": content},
            {"role": "user", "content": question}
        ]
        response = openai.ChatCompletion.create(model=model, messages=messages, temperature=0.1)
        return response['choices'][0]['message']['content']
    except Exception as e:
        return str(e)

def check_for_redundancy(text):
    question = f"Is the following content consistent and free of redundancies? If not, please refine it: '{text}'"
    response = get_answer("", question)
    return "yes" in response.lower()

def rephrase_to_remove_redundancy(text):
    question = f"Please rephrase the following content to remove redundancies: '{text}'"
    return get_answer("", question)


def remove_manual_renewal(text):
    potential_renewal_phrases = [
        "The parties may also renew this agreement in writing upon mutual Agreement",
        "The Parties may renew this Agreement in writing upon mutual Agreement"
        # You can add more potential phrases if needed
    ]

    for phrase in potential_renewal_phrases:
        text = text.replace(phrase, "")

    return text.strip()

def modify_term_of_agreement(paragraph, new_term, auto_renewal=False):
    modification_prompt = f"Modify the term of the agreement to {new_term} years. Original Content: '{paragraph.text}'"
    modification_response = get_answer(paragraph.text, modification_prompt)

    if auto_renewal:
        auto_renewal_clause = "This Agreement shall automatically renew for successive one (1) year terms unless either party provides written notice of its intention not to renew at least thirty (30) days prior to the end of the then-current term."
        modification_response += " " + auto_renewal_clause
        modification_response = remove_manual_renewal(modification_response)

    paragraph.clear()
    new_run = paragraph.add_run(modification_response)
    new_run.font.color.rgb = RGBColor(255, 0, 0)

def get_batch_answer(prompts_list):
    stringified_prompts = json.dumps(prompts_list)
    prompts = [{"role": "user", "content": stringified_prompts}]
    batch_instruction = {
        "role": "system",
        "content": "Complete every element of the array. Reply with an array of all completions."
    }
    prompts.append(batch_instruction)

    response = openai.ChatCompletion.create(model="gpt-4", messages=prompts, temperature=0.1)
    return json.loads(response.choices[0].message.content)

@app.route('/generate-agreement', methods=['POST'])
def generate_agreement():
    try:
        logger.info("Processing generate-agreement request.")

        data = request.json
        agreement_type = data.get('agreementType')
        parameters = data.get('parameters')

        template_paths = uploaded_templates.get(agreement_type)
        if not template_paths:
            return jsonify({"error": "Template not found!"}), 404

        modified_files = []

        for template_path in template_paths:
            last_processed_template = template_path  # Update the reference
            doc = Document(template_path)
            modifications_made = False

            for param in parameters:
                if param['key'].lower() in ['term of agreement', 'Term of agreement']:
                    for p in doc.paragraphs:
                        if p.text.startswith("This Agreement shall commence on the Effective Date above and shall terminate"):
                            modification_prompt = f"Modify the term of the agreement to {param['value']} years. Original Content: '{p.text}'"
                            modification_response = get_answer(p.text, modification_prompt)

                            sentences = sent_tokenize(modification_response)
                            if len(sentences) > 1:
                                refined_modification = ' '.join(sentences[:-1])  # join all but the last sentence
                            else:
                                refined_modification = modification_response 

                            p.clear()
                            new_run = p.add_run(refined_modification)
                            new_run.font.color.rgb = RGBColor(255, 0, 0)
                            modifications_made = True

                elif param['key'].lower() == 'exclusive manufacturer (section 3.1)':
                    for p in doc.paragraphs:
                        if p.text.startswith("3.1 Exclusive Territory Manufacturer"):
                            modification_prompt = f"Modify the exclusivity details for the manufacturer to have an exclusivity period of only one year and require a $100k payment to maintain exclusivity. Original Content: '{p.text}'"
                            modification_response = get_answer(p.text, modification_prompt)
                            p.clear()
                            new_run = p.add_run(modification_response)
                            new_run.font.color.rgb = RGBColor(255, 0, 0)
                            modifications_made = True
                            break

                else:
                    paragraphs = [p.text for p in doc.paragraphs]
                    relevance_prompts = [f"Does the paragraph specify the duration or termination criteria for the agreement? Content: '{para}'" for para in paragraphs]
                    with ProcessPoolExecutor() as executor:
                        batch_relevance_responses = list(executor.map(get_batch_answer, [relevance_prompts[i:i+10] for i in range(0, len(relevance_prompts), 10)]))
                    relevance_scores = [(para, score) for para, score in zip(paragraphs, [item for sublist in batch_relevance_responses for item in sublist]) if "yes" in score.lower()]
                    if not relevance_scores:
                        continue
                    most_relevant_paragraph = max(relevance_scores, key=lambda x: x[1])[0]
                    modification_prompt = f"Considering the context of the agreement, please suggest a modification to align with the concept of '{param['key']}' and incorporate the information: '{param['value']}'. Original Content: '{most_relevant_paragraph}'"
                    modification_response = get_answer(most_relevant_paragraph, modification_prompt)

                    if check_for_redundancy(modification_response):
                        modification_response = rephrase_to_remove_redundancy(modification_response)

                    for p in doc.paragraphs:
                        if most_relevant_paragraph in p.text:
                            p.clear()
                            old_run = p.add_run(most_relevant_paragraph)
                            modification_run = p.add_run(" [MODIFICATION: " + modification_response + "]")
                            modification_run.font.color.rgb = RGBColor(255, 0, 0)
                            modifications_made = True

            output_path = tempfile.mktemp(suffix=".docx")
            doc.save(output_path)
            if modifications_made:
                modified_files.append(output_path)

        if not modified_files:
            return jsonify({"message": "No modifications were made. Returning the original document."})


        if len(modified_files) == 1:
            directory, filename = os.path.split(modified_files[0])
            return send_from_directory(directory, filename, as_attachment=True, download_name="modified_agreement.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            # For simplicity, if both files are modified, we'll return a ZIP file containing both.
            zipname = tempfile.mktemp(suffix=".zip")
            with zipfile.ZipFile(zipname, 'w') as zf:
                for idx, f in enumerate(modified_files):
                    zf.write(f, f"modified_agreement_{idx+1}.docx")

            return send_file(zipname, as_attachment=True, download_name="modified_agreements.zip", mimetype='application/zip')

    except Exception as e:
        logger.error("Error encountered in generate-agreement: %s", str(e))
        logger.error(traceback.format_exc())  # Log the full traceback
        if last_processed_template:
            directory, filename = os.path.split(last_processed_template)
            return send_from_directory(directory, filename, as_attachment=True, download_name="original_agreement.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            return jsonify({"error": "An error occurred during processing."}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=int(os.environ.get('PORT', 8080)))